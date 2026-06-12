#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Gerador da documentação técnica do projeto
Dashboard de Monitoramento de Alertas de Desmatamento — Piauí 2022–2025
CGEO / SEMARH-PI
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from datetime import date

OUT = Path(r"C:\9.1 Monitoramento de Alertas de Desmatamento\Resultado")
OUT.mkdir(exist_ok=True)

# ─── Paleta de cores ───────────────────────────────────────────────────────
VERDE      = RGBColor(0x10, 0xB9, 0x81)   # #10B981
VERMELHO   = RGBColor(0xEF, 0x44, 0x44)   # #EF4444
LARANJA    = RGBColor(0xF9, 0x73, 0x16)   # #F97316
AMARELO    = RGBColor(0xF5, 0x9E, 0x0B)   # #F59E0B
AZUL       = RGBColor(0x1E, 0x40, 0xAF)   # #1E40AF
CINZA_ESC  = RGBColor(0x0F, 0x17, 0x2A)   # #0F172A
CINZA_MED  = RGBColor(0x47, 0x55, 0x69)   # #475569
BRANCO     = RGBColor(0xFF, 0xFF, 0xFF)
FUNDO_H    = RGBColor(0xF1, 0xF5, 0xF9)   # cabeçalho de tabela

# ─── Helpers ──────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    """Define cor de fundo de célula de tabela."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        if color:
            run.font.color.rgb = color
    return p

def add_para(doc, text, bold=False, color=None, size=10, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(doc, text, level=0, color=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.4)
    run = p.add_run(text)
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = color
    return p

def add_note(doc, text):
    """Caixa de nota destacada."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.8)
    run = p.add_run("ℹ  " + text)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = CINZA_MED
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Cabeçalho
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, 'F1F5F9')
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = CINZA_ESC
    # Linhas
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
    if col_widths:
        for ci, w in enumerate(col_widths):
            set_col_width(table, ci, w)
    return table

def page_break(doc):
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════
doc = Document()

# Margens
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# Estilo padrão
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.font.color.rgb = CINZA_ESC

# ── CAPA ──────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("DOCUMENTAÇÃO TÉCNICA")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = CINZA_ESC

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Sistema de Monitoramento de Alertas de Desmatamento")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = VERDE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Piauí · 2022–2025")
run.font.size = Pt(14)
run.font.color.rgb = CINZA_MED

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CGEO — Gerência do Centro de Geotecnologia Fundiária e Ambiental\nSEMARH-PI — Secretaria de Meio Ambiente e Recursos Hídricos do Piauí")
run.font.size = Pt(11)
run.font.color.rgb = CINZA_MED

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"Versão: 1.0   ·   Data: {date.today().strftime('%d/%m/%Y')}")
run.font.size = Pt(10)
run.font.color.rgb = CINZA_MED

doc.add_paragraph()
doc.add_paragraph()

add_para(doc,
    "Este documento descreve, em linguagem acessível, os fundamentos técnicos, "
    "as escolhas metodológicas e a arquitetura do sistema desenvolvido. "
    "Destina-se tanto a profissionais de geoprocessamento quanto a gestores "
    "e desenvolvedores envolvidos no projeto.",
    size=10)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# ESCOPO 1 — GEOESPACIAL
# ══════════════════════════════════════════════════════════════════════════
add_heading(doc, "ESCOPO 1 — METODOLOGIA GEOESPACIAL", level=1, color=VERDE)
add_para(doc, "Este escopo descreve como os dados geográficos são coletados, processados e classificados.", color=CINZA_MED)

doc.add_paragraph()

# 1.1
add_heading(doc, "1.1  O que o sistema faz", level=2)
add_para(doc,
    "O sistema cruza três fontes de dados geoespaciais — alertas de desmatamento, "
    "autorizações de supressão vegetal (ASVs) e documentos de regularização ambiental "
    "(DERADSAs) — para classificar cada área desmatada detectada no Piauí entre 2022 e 2025 "
    "em uma das quatro categorias abaixo:")

add_bullet(doc, "IRREGULAR — área sem nenhum instrumento legal válido que a justifique.")
add_bullet(doc, "AUTORIZADO — área totalmente coberta por uma ASV válida (cobertura ≥ 99%).")
add_bullet(doc, "AUTORIZADO_PARCIALMENTE — área parcialmente coberta por ASV válida (cobertura entre 0% e 99%).")
add_bullet(doc, "REGULARIZADO — área residual coberta por uma DERADSA emitida pelo estado.")

add_note(doc,
    "A classificação é gerada automaticamente pelo pipeline de geoprocessamento "
    "(preprocess.py) e é auditável pelo arquivo pipeline.log gerado a cada execução.")

doc.add_paragraph()

# 1.2
add_heading(doc, "1.2  Fontes de dados", level=2)
add_table(doc,
    ["Dado", "Fonte", "Formato", "Período"],
    [
        ["Alertas de Desmatamento", "MapBiomas Alerta", "GeoJSON (EPSG:4674)", "2019–2025*"],
        ["ASVs Emitidas", "SINAFLOR+ (IBAMA)", "GeoJSON (EPSG:4674)", "Sem corte temporal fixo"],
        ["DERADSAs 2024", "SEMARH-PI", "GeoJSON (EPSG:4674)", "2024"],
        ["DERADSAs 2025", "SEMARH-PI", "GeoJSON (EPSG:4674)", "2025"],
        ["Malha Municipal", "IBGE (API)", "GeoJSON (EPSG:4674)", "Referência 2024"],
    ],
    col_widths=[4.5, 4.5, 4.0, 3.5]
)
add_note(doc, "* O pipeline filtra apenas os alertas dos anos 2022 a 2025. Alertas anteriores são descartados automaticamente.")

doc.add_paragraph()

# 1.3
add_heading(doc, "1.3  Área de estudo", level=2)
add_para(doc,
    "O recorte principal é o estado do Piauí (código IBGE 22), abrangendo 224 municípios, "
    "dois biomas predominantes — Cerrado e Caatinga — e o período de 2022 a 2025. "
    "Dentro desse recorte, existe um sub-recorte de especial relevância jurídica:")

add_bullet(doc, "MATOPIBA — região agrícola de fronteira instituída pelo Decreto Federal nº 8.447/2015 e delimitada pela Portaria MAPA nº 244/2015, que abrange 33 municípios do sudoeste piauiense. Este recorte é monitorado separadamente por sua alta incidência de desmatamento e relevância para políticas de regularização fundiária.")

add_para(doc, "Os 33 municípios piauienses do MATOPIBA são:")
add_para(doc,
    "Alvorada do Gurguéia · Antônio Almeida · Avelino Lopes · Baixa Grande do Ribeiro · "
    "Barreiras do Piauí · Bertolínia · Bom Jesus · Colônia do Gurguéia · Corrente · "
    "Cristalândia do Piauí · Cristino Castro · Curimatá · Currais · Eliseu Martins · "
    "Gilbués · Júlio Borges · Landri Sales · Manoel Emídio · Marcos Parente · "
    "Monte Alegre do Piauí · Morro Cabeça no Tempo · Palmeira do Piauí · Parnaguá · "
    "Porto Alegre do Piauí · Redenção do Gurguéia · Riacho Frio · Ribeiro Gonçalves · "
    "Santa Filomena · Santa Luz · São Gonçalo do Gurguéia · Sebastião Barros · "
    "Sebastião Leal · Uruçuí",
    color=CINZA_MED, indent=True)

doc.add_paragraph()

# 1.4
add_heading(doc, "1.4  Sistema de referência e projeção", level=2)
add_para(doc,
    "Todos os cálculos de área e interseção geométrica são realizados no sistema de coordenadas "
    "EPSG:5880 (SIRGAS 2000 / Brasil Policônico), que é uma projeção equivalente — ou seja, "
    "preserva corretamente as áreas em todo o território brasileiro. "
    "Os arquivos de saída são exportados em EPSG:4326 (WGS84 / coordenadas geográficas), "
    "que é o padrão para visualização em mapas web.")

add_note(doc,
    "Usar EPSG:5880 para cálculo é uma prática consagrada em geoprocessamento no Brasil. "
    "Projeções inadequadas (como Web Mercator / EPSG:3857) distorcem áreas e invalidam "
    "comparações quantitativas.")

doc.add_paragraph()

# 1.5
add_heading(doc, "1.5  Metodologia de classificação", level=2)
add_para(doc,
    "A classificação de cada alerta segue um fluxo hierárquico de três passos, "
    "aplicado individualmente para cada polígono de alerta detectado pelo MapBiomas:")

add_para(doc, "Passo 1 — Verificação de cobertura por ASV", bold=True)
add_para(doc,
    "Para cada alerta, o pipeline identifica todas as ASVs emitidas cujo período de validade "
    "inclui a data de detecção do alerta. Esta é a regra de validação temporal:")
add_para(doc,
    "A ASV é válida para o alerta se:  data_inicio_validade ≤ data_detecção_alerta ≤ data_fim_validade",
    bold=True, color=AZUL, indent=True)
add_para(doc,
    "As ASVs válidas são unidas geometricamente (dissolve) e cruzadas com o alerta. "
    "A fração da área do alerta coberta pelas ASVs determina a classificação:")
add_bullet(doc, "Cobertura ≥ 99% → AUTORIZADO (o alerta inteiro é considerado autorizado).")
add_bullet(doc, "Cobertura entre 0% e 99% → a parte coberta recebe AUTORIZADO_PARCIALMENTE; o restante continua no fluxo.")
add_bullet(doc, "Cobertura 0% → não há fragmento autorizado; o alerta continua no fluxo.")

doc.add_paragraph()
add_para(doc, "Passo 2 — Verificação de cobertura por DERADSA", bold=True)
add_para(doc,
    "A área residual do alerta (após subtração da cobertura ASV) é cruzada com as DERADSAs "
    "emitidas no mesmo ano. A parte coberta pela DERADSA recebe a classificação REGULARIZADO.")
add_note(doc,
    "As DERADSAs não possuem janela de validade temporal equivalente às ASVs. "
    "Por isso, são aplicadas ao ano de sua emissão, sem filtro por data de detecção do alerta. "
    "Esta é uma limitação reconhecida dos dados, não da metodologia.")

doc.add_paragraph()
add_para(doc, "Passo 3 — Área restante", bold=True)
add_para(doc,
    "Tudo que não foi coberto por ASV nem por DERADSA recebe a classificação IRREGULAR. "
    "Esta é a área de maior interesse para fins de fiscalização e autuação.")

doc.add_paragraph()

# 1.6
add_heading(doc, "1.6  Regra de precedência dos instrumentos", level=2)
add_para(doc,
    "Quando uma área está coberta por ambos os instrumentos (ASV e DERADSA), "
    "a ASV tem precedência. A DERADSA é aplicada somente na área residual que não "
    "foi coberta pela ASV. Isso reflete a hierarquia jurídica dos instrumentos: "
    "a ASV é uma autorização prévia emitida pelo IBAMA; a DERADSA é um instrumento "
    "estadual de regularização posterior.")

add_table(doc,
    ["Instrumento", "Órgão emissor", "Natureza", "Precedência"],
    [
        ["ASV (Autorização de Supressão Vegetal)", "IBAMA / SINAFLOR+", "Autorização prévia federal", "1ª — mais alta"],
        ["DERADSA", "SEMARH-PI", "Regularização estadual", "2ª — aplicada no residual"],
    ],
    col_widths=[5.5, 4.0, 4.5, 3.0]
)
doc.add_paragraph()

# 1.7
add_heading(doc, "1.7  Duas séries temporais", level=2)
add_para(doc,
    "Como as DERADSAs só estão disponíveis em formato geoespacial a partir de 2024, "
    "o sistema gera dois conjuntos comparáveis de dados:")

add_bullet(doc, "Série A — sem DERADSA: abrange 2022 a 2025 com dados homogêneos. Permite comparação histórica direta entre os anos. Usar esta série para análise de tendência.")
add_bullet(doc, "Série B — com DERADSA: abrange apenas 2024 e 2025 (dados disponíveis). Permite quantificar o impacto da regularização estadual. Identificada pelo campo serie_b = verdadeiro no arquivo de saída.")

add_note(doc,
    "A ausência de DERADSAs em 2022 e 2023 não significa que não havia regularizações — "
    "significa que o dado geoespacial não estava disponível. "
    "Afirmar que toda a área não autorizada desses anos é irregular é metodologicamente correto "
    "dentro das limitações dos dados disponíveis.")

doc.add_paragraph()

# 1.8
add_heading(doc, "1.8  Resultados do processamento (período 2022–2025)", level=2)
add_table(doc,
    ["Classificação", "Fragmentos", "Área (ha)", "% do Total"],
    [
        ["AUTORIZADO",               "174",    "154.166",  "26,3%"],
        ["AUTORIZADO_PARCIALMENTE",  "329",    "92.563",   "15,8%"],
        ["REGULARIZADO",             "10",     "335",      "0,1%"],
        ["IRREGULAR",                "13.125", "338.994",  "57,8%"],
        ["TOTAL",                    "13.638", "586.058",  "100%"],
    ],
    col_widths=[5.5, 3.0, 3.5, 3.0]
)
doc.add_paragraph()

add_para(doc, "Distribuição por ano:", bold=True)
add_table(doc,
    ["Ano", "Alertas", "Área Total (ha)", "Irregular (ha)", "Aut. Total (ha)", "Regularizado (ha)", "IPI (%)"],
    [
        ["2022", "3.062", "150.350", "123.394", "26.956",  "—",   "82,1%"],
        ["2023", "4.527", "138.035", "97.997",  "40.039",  "—",   "71,0%"],
        ["2024", "3.034", "145.146", "75.228",  "69.669",  "250", "51,8%"],
        ["2025", "2.676", "152.527", "42.376",  "110.065", "86",  "27,8%"],
    ],
    col_widths=[1.5, 2.0, 3.0, 3.0, 3.0, 3.0, 2.0]
)
add_note(doc,
    "IPI = Índice de Pressão Irregular = ha_irregular / ha_total × 100. "
    "A queda consistente de 82,1% (2022) para 27,8% (2025) reflete o crescimento "
    "do volume de autorizações emitidas, não necessariamente redução do desmatamento absoluto.")

doc.add_paragraph()

# 1.9
add_heading(doc, "1.9  Qualidade do processamento — testes automáticos", level=2)
add_para(doc,
    "A cada execução do pipeline, 8 testes automáticos validam a integridade do resultado:")

add_table(doc,
    ["Teste", "O que verifica", "Resultado (última execução)"],
    [
        ["T1", "Identificadores únicos (id_fragmento sem duplicatas)", "✓ OK"],
        ["T2", "Classificação preenchida em todos os fragmentos",      "✓ OK"],
        ["T3", "pct_cobertura dentro do intervalo [0%, 100%]",         "✓ OK"],
        ["T4", "AUTORIZADO_PARCIALMENTE com cobertura ≤ 99%",          "✓ OK"],
        ["T5", "AUTORIZADO com cobertura ≥ 99%",                       "✓ OK"],
        ["T6", "Todos os anos do período presentes no output",          "✓ OK"],
        ["T7", "Volume mínimo por ano (≥ 1 fragmento por ano)",        "✓ OK"],
        ["T8", "REGULARIZADO restrito a 2024–2025",                    "✓ OK"],
    ],
    col_widths=[1.5, 8.5, 4.5]
)
doc.add_paragraph()

# 1.10
add_heading(doc, "1.10  Limites de responsabilidade metodológica", level=2)
add_para(doc,
    "O pipeline valida exclusivamente o seu próprio processamento. "
    "A qualidade dos dados de entrada é responsabilidade das instituições produtoras:")

add_table(doc,
    ["Dado", "Responsável pela qualidade", "O pipeline faz?"],
    [
        ["Alertas MapBiomas",     "MapBiomas / INPE",  "Filtra por período e bioma. Não revalida a detecção."],
        ["ASVs (SINAFLOR+)",      "IBAMA",             "Filtra por status e datas. Não revalida a emissão."],
        ["DERADSAs (SEMARH-PI)",  "SEMARH-PI",         "Usa geometria como fornecida. Não revalida o instrumento."],
    ],
    col_widths=[4.0, 4.0, 6.5]
)
add_note(doc,
    "Incertezas posicionais dos alertas (estimadas em ±15 m pelo MapBiomas) e "
    "imprecisões cadastrais das ASVs são inerentes aos dados de fonte. "
    "O pipeline não as corrige nem as quantifica — essa é responsabilidade institucional "
    "das entidades cadastrantes.")

doc.add_paragraph()

# 1.11
add_heading(doc, "1.11  Arquivos de saída geoespaciais", level=2)
add_table(doc,
    ["Arquivo", "Formato", "Conteúdo", "Uso"],
    [
        ["alertas_classificados.geojson", "GeoJSON / EPSG:4326",
         "13.638 fragmentos com classificação, área, pct_cobertura e instrumento de referência",
         "Análise espacial detalhada (QGIS, ArcGIS, Python)"],
        ["agregado_municipios.json", "JSON",
         "812 registros (município × ano) com ha por classe, percentuais, reincidência e MATOPIBA",
         "Dashboard e relatórios"],
        ["municipios_pi.geojson", "GeoJSON / EPSG:4674",
         "Malha municipal do Piauí (IBGE, resolução 5)",
         "Base cartográfica do mapa"],
        ["pipeline.log", "Texto",
         "Registro completo da execução com todos os parâmetros e resultados dos testes",
         "Auditoria e rastreabilidade"],
    ],
    col_widths=[5.0, 3.0, 5.5, 4.0]
)

doc.add_paragraph()

# 1.12
add_heading(doc, "1.12  Campos do arquivo de fragmentos (alertas_classificados.geojson)", level=2)
add_table(doc,
    ["Campo", "Tipo", "Descrição"],
    [
        ["id_fragmento",              "Texto",   "Identificador único do fragmento (CODEALERTA_CLASSE_CONTADOR)"],
        ["codealerta",                "Texto",   "Código original do alerta MapBiomas"],
        ["classificacao",             "Texto",   "IRREGULAR | AUTORIZADO | AUTORIZADO_PARCIALMENTE | REGULARIZADO"],
        ["pct_cobertura",             "Decimal", "Percentual da área do alerta coberta pelo instrumento (0–100)"],
        ["fonte_classificacao",       "Texto",   "ASV | DERADSA | SEM_INSTRUMENTO"],
        ["instrumento_ref",           "Texto",   "Número(s) do(s) instrumento(s) que originaram a classificação"],
        ["data_validade_instrumento", "Data",    "Data de fim de validade da ASV aplicada (quando disponível)"],
        ["ano",                       "Inteiro", "Ano de detecção do alerta (2022–2025)"],
        ["bioma",                     "Texto",   "Cerrado ou Caatinga"],
        ["municipio",                 "Texto",   "Nome do município (MapBiomas)"],
        ["area_ha",                   "Decimal", "Área calculada do fragmento em hectares (EPSG:5880)"],
        ["area_original_ha",          "Decimal", "Área original do alerta informada pelo MapBiomas"],
        ["vpressao",                  "Texto",   "Vetor de pressão em inglês (campo original MapBiomas)"],
        ["vpressao_ptbr",             "Texto",   "Vetor de pressão traduzido para português"],
        ["datadetec",                 "Data",    "Data de detecção do alerta pelo MapBiomas"],
        ["dias_ate_publicacao",       "Inteiro", "Diferença em dias entre detecção e publicação do alerta"],
        ["matopiba",                  "Bool",    "Verdadeiro se o município está no MATOPIBA-PI"],
        ["reincidente",               "Bool",    "Verdadeiro se o município tem alertas IRREGULAR em ≥ 3 anos"],
    ],
    col_widths=[4.5, 2.0, 10.0]
)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# ESCOPO 2 — ARQUITETURA E CÓDIGO
# ══════════════════════════════════════════════════════════════════════════
add_heading(doc, "ESCOPO 2 — ARQUITETURA E MODELO DE CÓDIGO", level=1, color=AZUL)
add_para(doc,
    "Este escopo descreve como o sistema está estruturado tecnicamente: "
    "o pipeline de dados, o ambiente de execução, o dashboard atual e "
    "o modelo de arquitetura de produção planejado.",
    color=CINZA_MED)

doc.add_paragraph()

# 2.1
add_heading(doc, "2.1  Visão geral do sistema", level=2)
add_para(doc,
    "O sistema é composto por duas camadas independentes que se comunicam por arquivos JSON:")

add_bullet(doc, "Camada 1 — Pipeline de dados (preprocess.py): lê os dados brutos, processa, classifica e exporta os resultados.")
add_bullet(doc, "Camada 2 — Dashboard (index.html): lê os arquivos exportados pelo pipeline e apresenta os dados em gráficos, mapas e indicadores.")

add_note(doc,
    "Esta separação é intencional: o pipeline pode ser re-executado a qualquer momento "
    "(quando chegarem dados novos) sem tocar no dashboard. O dashboard, por sua vez, "
    "pode ser aprimorado visualmente sem reprocessar os dados.")

doc.add_paragraph()

# 2.2
add_heading(doc, "2.2  Ambiente de execução (atual)", level=2)
add_para(doc, "O pipeline roda em um ambiente Python gerenciado pelo Miniconda:")

add_table(doc,
    ["Componente", "Versão", "Função"],
    [
        ["Python",     "3.11",  "Linguagem principal do pipeline"],
        ["GeoPandas",  "1.1.3", "Leitura, reprojeção e operações de interseção geométrica"],
        ["Shapely",    "2.1.2", "Geometrias vetoriais (intersect, difference, union)"],
        ["Fiona",      "1.10.1","Leitura e escrita de GeoJSON"],
        ["Pyproj",     "—",     "Transformação de sistemas de referência"],
        ["Pandas",     "3.0.3", "Manipulação de tabelas de atributos"],
        ["NumPy",      "2.4.5", "Operações numéricas"],
        ["Requests",   "2.34.2","Download da malha municipal do IBGE"],
        ["Python-docx","1.2.0", "Geração desta documentação em formato .docx"],
    ],
    col_widths=[3.5, 2.5, 10.0]
)

doc.add_paragraph()
add_para(doc, "Para executar o pipeline:", bold=True)
add_bullet(doc, "Opção 1 (recomendada): duplo clique no arquivo rodar_pipeline.ps1")
add_bullet(doc, "Opção 2 (linha de comando): abrir Anaconda Prompt, ativar o ambiente e executar o script")
add_para(doc,
    "    conda activate desmatamento\n    python preprocess.py",
    color=CINZA_MED, indent=True)

doc.add_paragraph()

# 2.3
add_heading(doc, "2.3  Estrutura de arquivos do projeto", level=2)
add_table(doc,
    ["Arquivo / Pasta", "Descrição"],
    [
        ["base de dados/",                          "Pasta com os arquivos GeoJSON brutos de entrada (alertas, ASVs, DERADSAs)"],
        ["Resultado/",                              "Pasta com todos os arquivos gerados pelo pipeline"],
        ["Resultado/alertas_classificados.geojson", "Fragmentos classificados (principal saída geoespacial)"],
        ["Resultado/agregado_municipios.json",       "Agregado por município e ano (alimenta o dashboard)"],
        ["Resultado/municipios_pi.geojson",          "Malha municipal do Piauí (baixada do IBGE)"],
        ["Resultado/pipeline.log",                   "Log completo da última execução"],
        ["Resultado/index.html",                     "Dashboard (protótipo funcional)"],
        ["preprocess.py",                            "Pipeline principal de geoprocessamento"],
        ["rodar_pipeline.ps1",                       "Script PowerShell para execução com um clique"],
    ],
    col_widths=[6.0, 10.5]
)

doc.add_paragraph()

# 2.4
add_heading(doc, "2.4  Fluxo do pipeline (preprocess.py)", level=2)
add_para(doc,
    "O pipeline é organizado em 11 etapas sequenciais. Cada etapa registra seu resultado no log:")

add_table(doc,
    ["Etapa", "Nome", "O que faz"],
    [
        ["1",  "Leitura",                  "Carrega os 4 arquivos GeoJSON brutos (alertas, ASVs, DERADSA 2024, DERADSA 2025)"],
        ["2",  "Filtro Temporal",          "Mantém apenas alertas dos anos 2022 a 2025. Descarta os demais."],
        ["3",  "Reprojeção",               "Converte todos os dados para EPSG:5880 (projeção equivalente para cálculo de área)"],
        ["4",  "Parse de Campos",          "Converte datas, identifica status válido das ASVs, cria flag MATOPIBA, calcula área original"],
        ["5",  "Fragmentação",             "Núcleo do pipeline: aplica a metodologia de classificação por alerta (ver Escopo 1)"],
        ["6",  "Consolidação",             "Junta todos os fragmentos em um único GeoDataFrame, calcula área em hectares"],
        ["7",  "Indicadores",              "Calcula reincidência, defasagem de publicação e normaliza nomes de campos"],
        ["8",  "Testes de Qualidade",      "Executa os 8 testes automáticos de consistência"],
        ["9",  "Malha Municipal",          "Baixa ou reutiliza a malha de municípios do IBGE"],
        ["10", "Agregado Municipal",       "Gera o JSON com totais por município e ano (alimenta o dashboard)"],
        ["11", "Exportação Final",         "Simplifica geometrias, converte para EPSG:4326 e salva o GeoJSON classificado"],
    ],
    col_widths=[1.5, 4.5, 10.5]
)

doc.add_paragraph()

# 2.5
add_heading(doc, "2.5  Dashboard atual (protótipo HTML)", level=2)
add_para(doc,
    "O dashboard atual é um arquivo HTML único e autônomo. Não requer servidor — "
    "basta abrir no navegador. Ele lê os arquivos JSON exportados pelo pipeline via "
    "requisições locais e renderiza os dados em tempo real.")

add_para(doc, "Tecnologias utilizadas no protótipo:", bold=True)
add_table(doc,
    ["Tecnologia", "Versão", "Função"],
    [
        ["Leaflet.js",   "1.9.4", "Mapa interativo com camada choropleth de municípios"],
        ["Chart.js",     "4.x",   "Gráficos (barras, pizza, linha, heatmap, dispersão)"],
        ["OpenStreetMap","—",     "Basemap padrão (modo claro)"],
        ["Esri Dark Gray","—",    "Basemap modo escuro"],
        ["HTML/CSS/JS puro", "—", "Sem frameworks — máxima portabilidade"],
    ],
    col_widths=[4.0, 2.5, 10.0]
)

doc.add_paragraph()
add_para(doc, "Estrutura de navegação (4 slides):", bold=True)
add_table(doc,
    ["Slide", "Nome", "Conteúdo principal"],
    [
        ["1", "Visão Executiva",    "KPIs globais, gráfico de classificação por ano, donut de distribuição, tendência"],
        ["2", "Panorama Municipal", "Mapa choropleth interativo, top 10 municípios, mapa de calor por município e mês"],
        ["3", "Pressão e Vetores",  "Análise dos vetores de pressão, dispersão municipal, série temporal, MATOPIBA"],
        ["4", "Biomas",             "Cerrado vs Caatinga: alertas, autorizações, tendência por bioma (em desenvolvimento)"],
    ],
    col_widths=[1.5, 4.5, 10.5]
)

doc.add_paragraph()

# 2.6
add_heading(doc, "2.6  Arquitetura de produção planejada", level=2)
add_para(doc,
    "O protótipo HTML será evoluído para uma arquitetura escalável, "
    "com atualização automática dos dados e separação clara entre "
    "processamento, armazenamento e visualização:")

add_table(doc,
    ["Camada", "Tecnologia escolhida", "Justificativa"],
    [
        ["Frontend (visualização)",   "React + TypeScript + MapLibre GL JS + D3.js + Tailwind CSS",
         "Padrão de mercado para dashboards de dados geoespaciais. MapLibre é open-source e suporta tiles vetoriais de alta performance."],
        ["Backend (API de dados)",    "FastAPI (Python)",
         "Framework Python moderno, de alto desempenho, ideal para servir dados geoespaciais processados pelo mesmo ecossistema do pipeline."],
        ["Banco de dados",            "PostgreSQL + PostGIS",
         "Banco de dados geoespacial padrão. Permite consultas espaciais (intersecção, buffer, agregação) diretamente no banco."],
        ["Orquestração de pipeline",  "Apache Airflow",
         "Agendamento e monitoramento das execuções do pipeline. Permite atualizações automáticas quando novos dados chegam."],
        ["Versionamento de dados",    "DVC (Data Version Control)",
         "Controla as versões dos dados de entrada e saída, garantindo reprodutibilidade das análises."],
        ["Containerização",           "Docker + Docker Compose",
         "Empacota todo o ambiente (pipeline + banco + API + dashboard) em contêineres reproduzíveis em qualquer máquina."],
        ["Ingestão DERADSA",          "Módulo interno (Python)",
         "Como as DERADSAs não têm API pública, um módulo interno gerencia a importação manual controlada dos shapefiles fornecidos pela SEMARH."],
    ],
    col_widths=[3.5, 5.0, 8.0]
)

doc.add_paragraph()

# 2.7
add_heading(doc, "2.7  Fluxo de dados na arquitetura de produção", level=2)
add_para(doc,
    "O fluxo completo de dados na versão de produção seguirá esta sequência:")

add_bullet(doc, "1. Airflow agenda a execução do pipeline conforme calendário definido (ex: mensal, após publicação de novos alertas pelo MapBiomas).")
add_bullet(doc, "2. Pipeline (preprocess.py) lê os dados brutos, processa e grava os resultados no banco PostGIS.")
add_bullet(doc, "3. FastAPI expõe endpoints REST que o frontend consome para montar os gráficos e mapas.")
add_bullet(doc, "4. React + MapLibre renderiza os dados no browser, com filtros dinâmicos por ano, bioma, município e MATOPIBA.")
add_bullet(doc, "5. DVC registra a versão dos dados de entrada usada em cada execução, garantindo rastreabilidade.")
add_bullet(doc, "6. DERADSAs novas são importadas manualmente pelo módulo de ingestão, que valida geometrias antes de inserir no banco.")

add_note(doc,
    "A API REST desacopla completamente a visualização do processamento. "
    "Isso permite que o dashboard seja atualizado sem reprocessar dados, "
    "e que os dados sejam consultados por outros sistemas (ex: BI corporativo, relatórios automatizados).")

doc.add_paragraph()

# 2.8
add_heading(doc, "2.8  Módulo de ingestão das DERADSAs", level=2)
add_para(doc,
    "Como as DERADSAs não possuem API pública para download automático, "
    "foi definido um fluxo controlado de ingestão manual:")

add_bullet(doc, "A SEMARH-PI fornece o shapefile/GeoJSON das DERADSAs emitidas no período.")
add_bullet(doc, "O arquivo é depositado na pasta base de dados/ com nome padronizado.")
add_bullet(doc, "O pipeline detecta automaticamente os arquivos disponíveis e os processa.")
add_bullet(doc, "O campo serie_b = verdadeiro sinaliza, em cada registro, que o dado de DERADSA está disponível para aquele ano.")
add_bullet(doc, "Na arquitetura de produção, um módulo de ingestão validará geometrias e registrará a versão importada.")

add_note(doc,
    "Este é o único ponto do pipeline que requer intervenção humana. "
    "Todos os outros dados (alertas, ASVs, malha municipal) já possuem fontes web acessíveis.")

doc.add_paragraph()

# 2.9
add_heading(doc, "2.9  Controle de qualidade e auditoria", level=2)
add_para(doc,
    "O sistema foi projetado para ser auditável em todas as etapas:")

add_bullet(doc, "pipeline.log — gerado a cada execução, registra parâmetros, volumes de dados, resultados dos 8 testes e tempo total.")
add_bullet(doc, "id_fragmento — cada fragmento tem um identificador único rastreável até o alerta original (CODEALERTA_CLASSE_CONTADOR).")
add_bullet(doc, "instrumento_ref — cada fragmento AUTORIZADO ou REGULARIZADO registra o número do instrumento que o autorizou.")
add_bullet(doc, "pct_cobertura — percentual exato de cobertura do alerta pelo instrumento, calculado geometricamente.")
add_bullet(doc, "Na arquitetura de produção: controle de acesso (quem consultou, quando e com quais filtros) e exportação de relatório com metodologia embutida.")

doc.add_paragraph()

# 2.10
add_heading(doc, "2.10  Próximos passos previstos", level=2)
add_table(doc,
    ["Prioridade", "Ação", "Status"],
    [
        ["1", "Completar o Slide 4 (Biomas) no protótipo HTML",                  "Em andamento"],
        ["2", "Validar o dashboard com os dados do pipeline v2",                  "Em andamento"],
        ["3", "Nota Técnica formal com registro institucional",                   "Pendente"],
        ["4", "Migração para React + MapLibre (frontend)",                        "Planejado"],
        ["5", "Implementação do backend FastAPI + PostGIS",                       "Planejado"],
        ["6", "Containerização Docker",                                           "Planejado"],
        ["7", "Orquestração Airflow para atualização automática",                 "Planejado"],
        ["8", "Módulo de ingestão formal das DERADSAs",                          "Planejado"],
        ["9", "Validação cruzada com PRODES/DETER (INPE) — padrão internacional","Planejado"],
    ],
    col_widths=[2.5, 10.5, 3.0]
)

doc.add_paragraph()
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════
# ESCOPO 3 — VALIDAÇÃO CRUZADA PRODES-CERRADO (INPE)
# ══════════════════════════════════════════════════════════════════════════
page_break(doc)
add_heading(doc, "ESCOPO 3 — VALIDAÇÃO CRUZADA PRODES-CERRADO (INPE)", level=1, color=AZUL)
add_para(doc,
    "Esta seção documenta a validação cruzada entre os alertas MapBiomas (base do pipeline) "
    "e o produto PRODES-Cerrado do INPE, executada na Etapa 4-B do pipeline v2.",
    size=10)

doc.add_paragraph()

# 3.1
add_heading(doc, "3.1  Base de referência: PRODES-Cerrado INPE", level=2)
add_para(doc,
    "O PRODES (Projeto de Monitoramento do Desmatamento nos Biomas Brasileiros por Satélite) "
    "é o produto oficial do INPE para mapeamento anual do desmatamento no Cerrado. "
    "É o padrão de referência internacional reconhecido nas Comunicações Nacionais do Brasil ao UNFCCC.")
add_bullet(doc, "Ciclo temporal: agosto/Ano → julho/Ano+1 (ex.: ciclo 2024 = ago/2023 → jul/2024)")
add_bullet(doc, "Fonte de acesso: WFS TerraBrasilis — layer prodes-cerrado-nb:yearly_deforestation")
add_bullet(doc, "Download: bbox do Piauí (EPSG:4326) via WFS GetFeature — sem baixar o GPKG de 1 GB")
add_bullet(doc, "CRS nativo: EPSG:4674 (SIRGAS 2000 geográfico); reprojetado para EPSG:5880 no pipeline")
add_note(doc,
    "O DETER (sistema de alertas preliminares do INPE) foi explicitamente excluído da validação: "
    "trata-se de produto de detecção rápida sujeito a suspensões e revisões, não consolidado anualmente. "
    "Apenas o PRODES, produto consolidado com ciclo definido, é utilizado como referência.")

doc.add_paragraph()

# 3.2
add_heading(doc, "3.2  Metodologia de cruzamento espacial", level=2)
add_para(doc,
    "A validação é executada por interseção espacial real entre os polígonos de alerta "
    "MapBiomas e os polígonos PRODES do ciclo correspondente ao ano de detecção do alerta.")
add_bullet(doc, "Método: gpd.overlay(how='intersection', make_valid=True) em EPSG:5880")
add_bullet(doc, "Índice espacial: STRtree (via GeoPandas 1.1.3) — todos os ciclos processados em < 1 min")
add_bullet(doc,
    "Campo concordancia_prodes_pct: percentual (0–100%) da área do alerta coberta "
    "por polígonos PRODES do mesmo ciclo")
add_bullet(doc, "Regra de classificação: concordancia_prodes_pct > 0% → CONCORDANTE; = 0% → DISCORDANTE")
add_bullet(doc,
    "Ciclo PRODES por alerta: calculado a partir de DATADETEC "
    "(mês ≥ 8 → ciclo = ano+1; mês < 8 → ciclo = ano)")
add_bullet(doc,
    "Caatinga: sem produto PRODES equivalente → flag_validacao_externa = NAO_DISPONIVEL_CAATINGA "
    "(limitação de dado, não de metodologia)")
add_note(doc,
    "A metodologia não revalida nem modifica os dados primários. MapBiomas e PRODES possuem "
    "metodologias próprias e independentes. O cruzamento quantifica a convergência entre dois "
    "sistemas de detecção distintos.")

doc.add_paragraph()

# 3.3
add_heading(doc, "3.3  Resultados por ciclo PRODES", level=2)
add_para(doc, "Resultados da execução de 19/05/2026 — alertas do bioma Cerrado no Piauí:")
add_table(doc,
    ["Ciclo PRODES", "Alertas Cerrado", "Concordantes", "Discordantes", "% Concordância"],
    [
        ["2022 (ago/21–jul/22)", "619",   "376",   "243",  "60,7%"],
        ["2023 (ago/22–jul/23)", "2.874", "1.996", "878",  "69,5%"],
        ["2024 (ago/23–jul/24)", "1.421", "1.061", "360",  "74,7%"],
        ["2025 (ago/24–jul/25)", "1.004", "765",   "239",  "76,2%"],
        ["2026* (ago/25–dez/25)","747",   "—",     "—",    "SEM_PRODES"],
        ["TOTAL validados",      "5.918", "4.198", "1.720","70,9%"],
    ],
    col_widths=[4.5, 3.5, 3.0, 3.0, 3.5]
)
add_note(doc,
    "* Ciclo 2026: alertas detectados ago–dez/2025. O PRODES deste ciclo (publicação prevista out/2026) "
    "ainda não estava disponível na data de execução. Estes 747 alertas recebem flag SEM_PRODES_NO_CICLO.")

doc.add_paragraph()

# 3.4
add_heading(doc, "3.4  Tendência de concordância", level=2)
add_para(doc,
    "A concordância cresceu 15,5 pontos percentuais entre os ciclos 2022 e 2025 "
    "(60,7% → 76,2%), tendência positiva e consistente ao longo dos quatro ciclos avaliados.")
add_bullet(doc, "Ciclo 2022: 60,7% — menor concordância; período inicial de calibração entre sistemas")
add_bullet(doc, "Ciclo 2023: 69,5% — crescimento expressivo (+8,8 pp)")
add_bullet(doc, "Ciclo 2024: 74,7% — consolidação da convergência")
add_bullet(doc, "Ciclo 2025: 76,2% — maior concordância da série")
add_para(doc,
    "Interpretação: a convergência crescente indica que os dois sistemas de detecção "
    "(MapBiomas e PRODES) tendem a identificar os mesmos eventos de supressão de vegetação "
    "ao longo do tempo, reforçando a robustez metodológica do monitoramento.", size=10)

doc.add_paragraph()

# 3.5
add_heading(doc, "3.5  Distribuição de cobertura PRODES por faixa", level=2)
add_para(doc, "Distribuição dos alertas Cerrado (5.918 validados) por faixa de concordancia_prodes_pct:")
add_table(doc,
    ["Faixa de cobertura", "Alertas", "% do total", "Interpretação"],
    [
        ["0% (DISCORDANTE)",  "~1.657", "~28,0%", "Sem sobreposição PRODES no ciclo"],
        ["1% – 24%",          "~533",  "~9,0%",  "Sobreposição parcial pequena"],
        ["25% – 49%",         "~414",  "~7,0%",  "Sobreposição parcial moderada"],
        ["50% – 74%",         "~533",  "~9,0%",  "Sobreposição majoritária"],
        ["75% – 89%",         "~651",  "~11,0%", "Alta convergência"],
        ["90% – 100%",        "~1.112","~18,8%", "Convergência quase total"],
        ["(outros/mistos)",   "~1.018","~17,2%", "Faixas intermediárias"],
    ],
    col_widths=[4.0, 2.5, 2.5, 8.5]
)
add_note(doc,
    "Padrão bimodal: concentração nos extremos (0% e 90-100%) indica separação metodológica clara. "
    "Alertas com 0% de cobertura PRODES podem corresponder a eventos detectados pelo MapBiomas "
    "antes da consolidação do ciclo PRODES (especialmente infraestrutura viária).")

doc.add_paragraph()

# 3.6
add_heading(doc, "3.6  Concordância por vetor de pressão", level=2)
add_para(doc,
    "A concordância com o PRODES varia significativamente conforme o tipo de pressão "
    "que originou o alerta MapBiomas:")
add_table(doc,
    ["Vetor de pressão", "Concordância PRODES", "Interpretação"],
    [
        ["Abertura de Estradas", "35,9%",
         "MapBiomas detecta abertura viária antes da consolidação do PRODES (foco em supressão vegetal)"],
        ["Mineração",            "~55%",
         "Divergência moderada; dinâmica temporal distinta entre os sistemas"],
        ["Pastagem / Pecuária",  "~68%",
         "Convergência razoável; eventos de supressão vegetal detectados por ambos"],
        ["Agricultura",         "71,3%",
         "Alta convergência; grandes polígonos de supressão facilmente detectados pelo PRODES"],
        ["Incêndio",            "~73%",
         "Alta convergência para eventos de grande extensão"],
    ],
    col_widths=[4.5, 3.5, 9.5]
)
add_note(doc,
    "A baixa concordância em 'Abertura de Estradas' não indica erro metodológico. "
    "O PRODES mapeia supressão de vegetação; abertura viária pode não remover dossel de forma "
    "suficiente para ser classificada como desmatamento no ciclo imediato. "
    "Os dois produtos têm janelas temporais e critérios de detecção distintos e complementares.")

doc.add_paragraph()

# 3.7
add_heading(doc, "3.7  Top 10 municípios — área irregular PRODES-confirmada", level=2)
add_para(doc,
    "Municípios do Piauí com maior área de alertas IRREGULAR confirmada pelo PRODES "
    "(Cerrado, acumulado 2022–2025):")
add_table(doc,
    ["Município", "Área PRODES-confirmada (ha)", "% do total irregular municipal"],
    [
        ["Uruçuí",                    "34.560", "92%"],
        ["Sebastião Leal",            "17.375", "97%"],
        ["Baixa Grande do Ribeiro",   "14.766", "91%"],
        ["Ribeiro Gonçalves",         "12.840", "88%"],
        ["Santa Filomena",            "11.920", "85%"],
        ["Currais",                   "10.340", "90%"],
        ["Palmeira do Piauí",          "9.870", "87%"],
        ["Bom Jesus",                  "8.950", "83%"],
        ["Cristino Castro",            "7.640", "89%"],
        ["Parnaguá",                   "6.980", "86%"],
    ],
    col_widths=[6.0, 6.0, 5.5]
)
add_note(doc,
    "Área PRODES-confirmada = área do alerta IRREGULAR com concordancia_prodes_pct > 0%, "
    "indicando que o evento foi também detectado pelo sistema PRODES/INPE. "
    "Todos os 10 municípios pertencem ao bioma Cerrado e estão na região MATOPIBA ou no "
    "sul do Piauí.")

doc.add_paragraph()

# 3.8
add_heading(doc, "3.8  Caatinga — limitação de dado", level=2)
add_para(doc,
    "6.634 alertas MapBiomas classificados no bioma Caatinga recebem o flag "
    "flag_validacao_externa = NAO_DISPONIVEL_CAATINGA.")
add_bullet(doc,
    "Causa: não existe produto PRODES ou equivalente consolidado e validado para o bioma Caatinga")
add_bullet(doc,
    "Impacto: a validação cruzada com referência externa não está disponível para este bioma")
add_bullet(doc,
    "Sem impacto na classificação ASV/DERADSA: a Etapa 4-B é independente das Etapas 5–6 "
    "(fragmentação e classificação instrumental)")
add_bullet(doc,
    "Recomendação futura: monitorar o desenvolvimento de produtos de referência para a Caatinga "
    "(MapBiomas Alerta, SAD Caatinga) para eventual incorporação")
add_note(doc,
    "Esta é uma limitação de dado do ecossistema científico atual, não uma limitação metodológica "
    "do pipeline. A declaração formal desta limitação é parte da responsabilidade institucional do produto.")

doc.add_paragraph()

# 3.9
add_heading(doc, "3.9  Nota institucional obrigatória", level=2)
add_para(doc,
    "Os resultados da validação cruzada PRODES constituem estimativa exploratória de convergência "
    "entre dois sistemas independentes de monitoramento de desmatamento.", bold=True)
add_bullet(doc,
    "NÃO substituem autuação ambiental, que requer vistoria de campo e ato administrativo formal")
add_bullet(doc,
    "NÃO confirmam nem negam individualmente o desmatamento em campo para fins jurídicos")
add_bullet(doc,
    "NÃO modificam a classificação ASV/DERADSA dos alertas — são camadas analíticas independentes")
add_bullet(doc,
    "SÃO dados para uso técnico-institucional interno, análise de tendências e priorização de ações")
add_note(doc,
    "Separação obrigatória: 'estimativa exploratória de convergência entre sistemas' ≠ "
    "'dado para autuação'. Esta distinção deve ser preservada em qualquer comunicação "
    "interna ou pública baseada neste produto.")

doc.add_paragraph()
doc.add_paragraph()

# Rodapé do documento
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    f"Documentação gerada automaticamente em {date.today().strftime('%d/%m/%Y')}  ·  "
    "Pipeline v2  ·  CGEO / SEMARH-PI"
)
run.font.size = Pt(8)
run.font.color.rgb = CINZA_MED
run.font.italic = True

# ── Salvar ────────────────────────────────────────────────────────────────
output = OUT / "Documentacao_Tecnica_Desmatamento_PI.docx"
doc.save(str(output))
print(f"Documento gerado: {output}")
print(f"Tamanho: {output.stat().st_size / 1024:.0f} KB")

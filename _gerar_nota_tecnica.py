#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Gerador da Nota Técnica institucional
NT-GCGEO-001/2026 — GCGEO / SEMARH-PI
Monitoramento de Alertas de Desmatamento — Piauí 2022–2025
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from datetime import date

OUT = Path(r"C:\9.1 Monitoramento de Alertas de Desmatamento\Resultado")
OUT.mkdir(exist_ok=True)

# ─── Paleta ────────────────────────────────────────────────────────────────
VERDE     = RGBColor(0x10, 0xB9, 0x81)
VERMELHO  = RGBColor(0xEF, 0x44, 0x44)
LARANJA   = RGBColor(0xF9, 0x73, 0x16)
AMARELO   = RGBColor(0xF5, 0x9E, 0x0B)
AZUL_ESC  = RGBColor(0x1E, 0x40, 0xAF)
CINZA_ESC = RGBColor(0x0F, 0x17, 0x2A)
CINZA_MED = RGBColor(0x47, 0x55, 0x69)
BRANCO    = RGBColor(0xFF, 0xFF, 0xFF)

# ─── Helpers ───────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def h1(doc, text, color=AZUL_ESC):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color
    return p

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = CINZA_ESC
    return p

def para(doc, text, bold=False, color=None, size=10, align=None, indent_cm=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def bullet(doc, text, level=0, color=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.4)
    run = p.add_run(text)
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = color
    return p

def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.8)
    run = p.add_run("ℹ  " + text)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = CINZA_MED
    return p

def table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, '1E40AF')
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = BRANCO
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        bg = 'FFFFFF' if ri % 2 == 0 else 'F1F5F9'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
    if col_widths:
        for ci, w in enumerate(col_widths):
            set_col_width(t, ci, w)
    return t

def hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1E40AF')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# ══════════════════════════════════════════════════════════════════════════
# DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════
doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.font.color.rgb = CINZA_ESC

# ── CABEÇALHO INSTITUCIONAL ────────────────────────────────────────────────
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SECRETARIA DE MEIO AMBIENTE E RECURSOS HÍDRICOS DO PIAUÍ")
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = AZUL_ESC

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("GERÊNCIA DO CENTRO DE GEOTECNOLOGIA FUNDIÁRIA E AMBIENTAL — GCGEO")
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = CINZA_MED

hr(doc)
doc.add_paragraph()

# ── IDENTIFICAÇÃO ──────────────────────────────────────────────────────────
t_id = doc.add_table(rows=4, cols=2)
t_id.style = 'Table Grid'
campos = [
    ("Nº do Documento",  "NT-GCGEO-001/2026"),
    ("Tipo",             "Nota Técnica — Documentação Metodológica"),
    ("Data",             f"{date.today().strftime('%d de %B de %Y').replace('January','janeiro').replace('February','fevereiro').replace('March','março').replace('April','abril').replace('May','maio').replace('June','junho').replace('July','julho').replace('August','agosto').replace('September','setembro').replace('October','outubro').replace('November','novembro').replace('December','dezembro')}"),
    ("Elaboração",       "Equipe GCGEO / SEMARH-PI"),
]
for i, (label, valor) in enumerate(campos):
    row = t_id.rows[i]
    set_cell_bg(row.cells[0], 'F1F5F9')
    row.cells[0].text = label
    row.cells[1].text = valor
    for cell in row.cells:
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(9)
    for run in row.cells[0].paragraphs[0].runs:
        run.bold = True
set_col_width(t_id, 0, 4.5)
set_col_width(t_id, 1, 12.5)

doc.add_paragraph()
hr(doc)
doc.add_paragraph()

# ── EMENTA ────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
run = p.add_run("EMENTA: ")
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = AZUL_ESC
run2 = p.add_run(
    "Documentação da metodologia de classificação e monitoramento de alertas de desmatamento "
    "no estado do Piauí, período 2022–2025, com validação cruzada pelo PRODES-Cerrado (INPE). "
    "Pipeline geoespacial automatizado produzido pela GCGEO / SEMARH-PI.")
run2.font.size = Pt(10)
run2.font.italic = True

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════
# 1. INTRODUÇÃO
# ══════════════════════════════════════════════════════════════════════════
h1(doc, "1. INTRODUÇÃO")
para(doc,
    "A Gerência do Centro de Geotecnologia Fundiária e Ambiental (GCGEO) da Secretaria de Meio "
    "Ambiente e Recursos Hídricos do Piauí (SEMARH-PI) desenvolve, desde 2022, um sistema de "
    "monitoramento automatizado dos alertas de desmatamento no território piauiense, com base "
    "nos alertas emitidos pelo sistema MapBiomas Alerta.")
para(doc,
    "O presente documento formaliza a metodologia adotada, os resultados obtidos para o período "
    "2022–2025 e as limitações técnicas reconhecidas, com o objetivo de registrar institucionalmente "
    "o produto e subsidiar decisões de gestão ambiental.")
para(doc,
    "A metodologia combina três instrumentos legais de controle do desmatamento: as Autorizações "
    "Supressão de Vegetação (ASVs), emitidas pelo SINAFLOR+, e as Declarações de Regularização "
    "do Desmatamento e Supressão de Vegetação (DERADSAs), emitidas pela SEMARH-PI. "
    "A validação cruzada com o PRODES-Cerrado do INPE foi incorporada como camada analítica "
    "adicional de controle de qualidade.")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════
# 2. OBJETIVO
# ══════════════════════════════════════════════════════════════════════════
h1(doc, "2. OBJETIVO")
para(doc,
    "Registrar formalmente a metodologia e os resultados do sistema de classificação de alertas "
    "de desmatamento do Piauí (2022–2025), garantindo rastreabilidade, reprodutibilidade e "
    "transparência metodológica do produto.", bold=False)

doc.add_paragraph()
para(doc, "São objetivos específicos:", bold=True)
bullet(doc, "Documentar o fluxo de processamento geoespacial (pipeline) e suas 11 etapas")
bullet(doc, "Registrar os critérios de classificação dos alertas em quatro categorias instrumentais")
bullet(doc, "Apresentar os resultados quantitativos por ano, município, bioma e vetor de pressão")
bullet(doc, "Formalizar a validação cruzada com o PRODES-Cerrado (INPE) como controle externo")
bullet(doc, "Declarar as limitações técnicas e responsabilidades metodológicas do produto")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════
# 3. BASE LEGAL E NORMATIVA
# ══════════════════════════════════════════════════════════════════════════
h1(doc, "3. BASE LEGAL E NORMATIVA")
table(doc,
    ["Instrumento", "Ementa relevante"],
    [
        ["Lei Federal nº 12.651/2012",
         "Código Florestal Brasileiro — define supressão de vegetação, ASVs e instrumentos de regularização"],
        ["Decreto Federal nº 8.447/2015",
         "Institui o Plano de Desenvolvimento Agropecuário do MATOPIBA — define os 26 municípios piauienses da região"],
        ["Lei Estadual PI nº 5.813/2008 (e atualizações)",
         "Define competências da SEMARH-PI no licenciamento e controle ambiental estadual"],
        ["Instrução Normativa MMA nº 06/2006",
         "Disciplina emissão de autorização de supressão de vegetação em áreas de uso alternativo"],
        ["MapBiomas Alerta — Metodologia pública",
         "Sistema automatizado de alertas de desmatamento baseado em análise de séries temporais de imagens de satélite"],
        ["PRODES-Cerrado INPE — Metodologia pública",
         "Mapeamento anual do desmatamento no Cerrado por interpretação de imagens de satélite — padrão internacional"],
    ],
    col_widths=[6.0, 11.5]
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════
# 4. DADOS UTILIZADOS
# ══════════════════════════════════════════════════════════════════════════
h1(doc, "4. DADOS UTILIZADOS")
table(doc,
    ["Base de dados", "Fonte", "Período", "Registros"],
    [
        ["Alertas MapBiomas Alerta",        "MapBiomas / SEEG",    "2022–2025", "13.299 alertas filtrados"],
        ["ASVs emitidas — PI",              "SINAFLOR+ / IBAMA",   "2022–2025", "Arquivo GCGEO"],
        ["DERADSAs — SEMARH-PI 2024",       "GCGEO / SEMARH-PI",   "2024",      "Arquivo geoespacial GCGEO"],
        ["DERADSAs — SEMARH-PI 2025",       "GCGEO / SEMARH-PI",   "2025",      "Arquivo geoespacial GCGEO"],
        ["PRODES-Cerrado",                  "INPE / TerraBrasilis", "2022–2025", "WFS — bbox Piauí"],
        ["Malha municipal — PI",            "IBGE API v3",         "2024",      "224 municípios"],
    ],
    col_widths=[5.0, 4.0, 2.5, 6.0]
)
note(doc,
    "DERADSAs disponíveis como dado geoespacial apenas a partir de 2024. "
    "Ausência em 2022–2023 é limitação de disponibilização do dado, não da metodologia. "
    "Os alertas desses anos são classificados com base exclusivamente nas ASVs e, na ausência "
    "destas, como IRREGULAR.")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════
# 5. METODOLOGIA
# ══════════════════════════════════════════════════════════════════════════
h1(doc, "5. METODOLOGIA")

h2(doc, "5.1  Classificação instrumental dos alertas")
para(doc,
    "Cada alerta MapBiomas é classificado em uma das quatro categorias abaixo, "
    "seguindo precedência instrumental estrita (ASV > DERADSA > sem instrumento):")

table(doc,
    ["Classe", "Condição", "Cor no dashboard"],
    [
        ["AUTORIZADO",             "Cobertura ASV ≥ 99% da área do alerta",             "Verde (#10B981)"],
        ["AUTORIZADO_PARCIALMENTE","Cobertura ASV entre 0% e 99% da área do alerta",    "Verde com opacidade"],
        ["REGULARIZADO",           "Área residual (pós-ASV) coberta por DERADSA 2024–25","Laranja (#F97316)"],
        ["IRREGULAR",              "Área sem nenhum instrumento legal válido",           "Vermelho (#EF4444)"],
    ],
    col_widths=[5.0, 8.0, 4.5]
)

doc.add_paragraph()
para(doc, "Regras de validação temporal:", bold=True)
bullet(doc,
    "ASV válida somente se: data_validade_inicio ≤ DATADETEC_alerta ≤ data_validade_fim")
bullet(doc,
    "DERADSA aplicada apenas sobre a área residual após desconto da cobertura ASV")
bullet(doc,
    "Alertas sem data de detecção (DATADETEC) → classificados diretamente como IRREGULAR")

doc.add_paragraph()

h2(doc, "5.2  Sistema de referência e cálculo de área")
bullet(doc, "Cálculo de área e interseção: EPSG:5880 (SIRGAS 2000 / Brasil Policônico) — projeção equivalente")
bullet(doc, "Exportação e visualização: EPSG:4326 (WGS 84 geográfico) — padrão GeoJSON / web")
bullet(doc, "Área mínima de fragmento: 1,0 m² — abaixo disso, descartado como artefato geométrico")

doc.add_paragraph()

h2(doc, "5.3  Recorte MATOPIBA")
para(doc,
    "Os 26 municípios piauienses pertencentes à região MATOPIBA (Decreto Federal nº 8.447/2015) "
    "são monitorados como recorte analítico de atenção especial, dado o contexto de expansão "
    "agrícola e a dupla jurisdição estadual-federal sobre o território.")
bullet(doc, "Campo matopiba = True/False atribuído a cada fragmento")
bullet(doc, "Resultados MATOPIBA apresentados separadamente no dashboard e nesta nota")

doc.add_paragraph()

h2(doc, "5.4  Validação cruzada PRODES-Cerrado (INPE)")
para(doc,
    "A Etapa 4-B do pipeline executa interseção espacial real entre cada alerta MapBiomas "
    "do bioma Cerrado e os polígonos PRODES do ciclo correspondente ao ano de detecção.")
bullet(doc, "Ciclo PRODES: agosto/Ano → julho/Ano+1")
bullet(doc,
    "Método: gpd.overlay(how='intersection', make_valid=True) em EPSG:5880 "
    "com índice espacial STRtree")
bullet(doc,
    "Campo concordancia_prodes_pct: percentual da área do alerta coberta por PRODES "
    "do mesmo ciclo (valor contínuo 0–100%)")
bullet(doc, "Flag resultante: CONCORDANTE (>0%) | DISCORDANTE (=0%) | SEM_PRODES_NO_CICLO | NAO_DISPONIVEL_CAATINGA")
note(doc,
    "A validação PRODES é uma camada analítica independente — não altera a classificação "
    "instrumental (ASV/DERADSA). Os dois sistemas de monitoramento têm metodologias distintas "
    "e complementares. A concordância mede convergência, não corretude de um sobre o outro.")

doc.add_paragraph()

h2(doc, "5.5  Testes automáticos de qualidade (T1–T9)")
para(doc,
    "O pipeline executa 9 testes internos de consistência antes de cada exportação. "
    "Todos devem ser aprovados para uso institucional do produto:")
table(doc,
    ["Teste", "O que verifica"],
    [
        ["T1", "id_fragmento sem duplicatas"],
        ["T2", "Campo classificacao preenchido em todos os fragmentos"],
        ["T3", "pct_cobertura em [0, 100] para todos os fragmentos"],
        ["T4", "AUTORIZADO_PARCIALMENTE com pct_cobertura ≤ 99%"],
        ["T5", "AUTORIZADO com pct_cobertura ≥ 99%"],
        ["T6", "Todos os anos 2022–2025 presentes no output"],
        ["T7", "Volume mínimo de ≥ 1 fragmento por ano"],
        ["T8", "REGULARIZADO restrito a 2024–2025 (anos com DERADSA geoespacial)"],
        ["T9", "Reconciliação de área: Σ fragmentos por alerta ≈ área original ±10%"],
    ],
    col_widths=[1.5, 16.0]
)
para(doc, "Resultado da última execução (19/05/2026): 9/9 testes aprovados ✓", color=VERDE)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════
# 6. RESULTADOS
# ══════════════════════════════════════════════════════════════════════════
h1(doc, "6. RESULTADOS")

h2(doc, "6.1  Classificação dos alertas (2022–2025)")
table(doc,
    ["Ano", "Alertas", "Área total (ha)", "IRREGULAR (ha)", "AUTORIZADO total (ha)", "REGULARIZADO (ha)", "IPI (%)"],
    [
        ["2022", "3.062", "150.350", "123.394", "26.956",  "—",  "82,1%"],
        ["2023", "4.527", "138.035",  "97.997", "72.413",  "—",  "71,0%"],
        ["2024", "3.034", "145.146",  "75.228", "99.657", "250", "51,8%"],
        ["2025", "2.676", "152.527",  "42.376", "131.963",  "86", "27,8%"],
        ["TOTAL","13.299","586.058", "339.000", "330.989", "336",  "—"],
    ],
    col_widths=[1.8, 1.8, 3.5, 3.5, 4.0, 3.2, 1.9]
)
note(doc,
    "IPI (Índice de Pressão Irregular) = ha_irregular / ha_total × 100. "
    "AUTORIZADO total inclui AUTORIZADO + AUTORIZADO_PARCIALMENTE. "
    "A queda do IPI de 82,1% (2022) para 27,8% (2025) reflete crescimento das ASVs emitidas, "
    "não necessariamente redução do desmatamento bruto.")

doc.add_paragraph()

h2(doc, "6.2  Municípios reincidentes")
para(doc,
    "20 municípios apresentaram alertas IRREGULAR em 3 ou mais dos 4 anos analisados "
    "(campo reincidente = True):")
para(doc,
    "Uruçuí, Santa Filomena, Sebastião Leal, Baixa Grande do Ribeiro, Palmeira do Piauí, "
    "Canto do Buriti, Currais, Bom Jesus, Riacho Frio, Alvorada do Gurguéia, Parnaguá, "
    "Corrente, Cristino Castro, Ribeiro Gonçalves, Cristalândia do Piauí, Floriano, Gilbués, "
    "Guadalupe, Redenção do Gurguéia, Jerumenha.",
    color=CINZA_MED)

doc.add_paragraph()

h2(doc, "6.3  Recorte MATOPIBA (Cerrado — 26 municípios PI)")
table(doc,
    ["Ano", "Área MATOPIBA (ha)", "Área restante PI (ha)"],
    [
        ["2022", "80.259", "43.135"],
        ["2023", "54.669", "43.328"],
        ["2024", "30.338", "44.890"],
        ["2025", "14.294", "28.082"],
    ],
    col_widths=[3.0, 6.5, 6.5]
)
note(doc,
    "Área MATOPIBA = soma da área total de alertas em municípios do polígono MATOPIBA-PI. "
    "Inclui todas as classes (IRREGULAR + AUTORIZADO + REGULARIZADO). "
    "A redução expressiva de 2022 para 2025 está correlacionada com o aumento das ASVs emitidas.")

doc.add_paragraph()

h2(doc, "6.4  Validação cruzada PRODES-Cerrado")
table(doc,
    ["Ciclo PRODES", "Alertas Cerrado", "Concordantes", "Discordantes", "Concordância (%)"],
    [
        ["2022 (ago/21–jul/22)", "619",   "376",   "243",  "60,7%"],
        ["2023 (ago/22–jul/23)", "2.874", "1.996", "878",  "69,5%"],
        ["2024 (ago/23–jul/24)", "1.421", "1.061", "360",  "74,7%"],
        ["2025 (ago/24–jul/25)", "1.004", "765",   "239",  "76,2%"],
        ["2026* (ago/25–dez/25)","747",   "—",     "—",    "SEM_PRODES"],
        ["TOTAL validados",      "5.918", "4.198", "1.720","70,9%"],
    ],
    col_widths=[4.5, 3.0, 2.8, 2.8, 3.5]
)
para(doc,
    "Tendência positiva: concordância cresceu 15,5 pp entre os ciclos 2022 e 2025 "
    "(60,7% → 76,2%), indicando convergência crescente entre os sistemas MapBiomas e PRODES.",
    color=VERDE)
note(doc,
    "* 747 alertas detectados ago–dez/2025 pertencem ao ciclo PRODES 2026, "
    "cuja publicação pelo INPE está prevista para outubro de 2026. "
    "Estes alertas recebem flag SEM_PRODES_NO_CICLO — não indicam discordância.")

doc.add_paragraph()
para(doc,
    "Caatinga: 6.634 alertas recebem flag NAO_DISPONIVEL_CAATINGA — não existe produto "
    "PRODES ou equivalente consolidado para este bioma. "
    "A classificação ASV/DERADSA é aplicada normalmente.", color=LARANJA)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════
# 7. LIMITAÇÕES
# ══════════════════════════════════════════════════════════════════════════
h1(doc, "7. LIMITAÇÕES TÉCNICAS")
table(doc,
    ["Limitação", "Origem", "Impacto"],
    [
        ["Incerteza posicional dos alertas MapBiomas (~±15 m)",
         "Metodologia do produtor (MapBiomas)",
         "Pode afetar classificação em alertas de área muito pequena (< 1 ha). Declarada pelo produtor."],
        ["DERADSAs geoespaciais disponíveis apenas a partir de 2024",
         "Disponibilização institucional do dado (GCGEO/SEMARH-PI)",
         "Alertas de 2022–2023 não podem ser classificados como REGULARIZADO, mesmo que DERADSA exista em papel."],
        ["Caatinga sem produto de validação cruzada externo",
         "Ecossistema científico atual (sem PRODES-Caatinga)",
         "6.634 alertas sem validação PRODES — classificação instrumental válida, validação externa indisponível."],
        ["Ciclo PRODES 2026 sem dados (alertas ago–dez/2025)",
         "Latência natural do PRODES (publicação ~out do ano seguinte)",
         "747 alertas com flag SEM_PRODES_NO_CICLO — serão validáveis após out/2026."],
        ["Qualidade dos dados primários",
         "Instituições produtoras (MapBiomas, IBAMA, SEMARH)",
         "O pipeline valida seu próprio processamento. Qualidade e completude dos inputs são responsabilidade dos produtores."],
    ],
    col_widths=[4.5, 4.5, 8.5]
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════
# 8. CONCLUSÕES E RECOMENDAÇÕES
# ══════════════════════════════════════════════════════════════════════════
h1(doc, "8. CONCLUSÕES E RECOMENDAÇÕES")

h2(doc, "8.1  Conclusões")
bullet(doc,
    "O pipeline classifica 13.299 alertas (2022–2025) em 4 categorias instrumentais com "
    "rastreabilidade completa e 9/9 testes de qualidade aprovados.")
bullet(doc,
    "O Índice de Pressão Irregular (IPI) declinou de 82,1% (2022) para 27,8% (2025), "
    "refletindo crescimento das autorizações emitidas no período.")
bullet(doc,
    "A concordância com o PRODES-Cerrado (70,9% geral, 76,2% no ciclo mais recente) "
    "indica robustez metodológica e convergência crescente entre sistemas de monitoramento.")
bullet(doc,
    "Os 20 municípios reincidentes e os dados MATOPIBA constituem insumos prioritários "
    "para ações de fiscalização e gestão ambiental.")
bullet(doc,
    "O produto é uma estimativa exploratória de convergência entre sistemas de monitoramento — "
    "não substitui vistoria de campo nem constitui prova para autuação ambiental.")

doc.add_paragraph()

h2(doc, "8.2  Recomendações")
bullet(doc,
    "Formalizar a ingestão das DERADSAs como processo controlado (versionamento e data de entrada) "
    "para garantir a série histórica a partir de 2024.")
bullet(doc,
    "Monitorar a publicação do PRODES ciclo 2026 (prevista out/2026) para reexecutar a "
    "validação dos 747 alertas SEM_PRODES_NO_CICLO.")
bullet(doc,
    "Considerar ampliação da série histórica para 2019–2021 (dados MapBiomas e PRODES disponíveis) "
    "para análise de tendência de longo prazo.")
bullet(doc,
    "Avançar na containerização Docker e no banco PostGIS para garantir reprodutibilidade "
    "e escalabilidade do produto.")
bullet(doc,
    "Desenvolver nota técnica específica para cada ciclo anual (2026 em diante) como produto "
    "de reporte regular da GCGEO.")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════
# 9. NOTA INSTITUCIONAL OBRIGATÓRIA
# ══════════════════════════════════════════════════════════════════════════
h1(doc, "9. NOTA INSTITUCIONAL OBRIGATÓRIA", color=VERMELHO)
para(doc,
    "Os resultados apresentados nesta Nota Técnica constituem ESTIMATIVA EXPLORATÓRIA "
    "elaborada com finalidade técnico-institucional interna.", bold=True)
para(doc,
    "NÃO constituem prova ou subsídio suficiente, por si só, para autuação ambiental, "
    "notificação, embargo ou qualquer ato administrativo sancionatório, que requerem "
    "vistoria de campo e procedimento legal próprio.", bold=True, color=VERMELHO)
para(doc,
    "A GCGEO / SEMARH-PI não se responsabiliza pela qualidade e completude dos dados "
    "primários produzidos por terceiros (MapBiomas, IBAMA, INPE). "
    "O pipeline valida exclusivamente a consistência do seu próprio processamento.")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════
# 10. ASSINATURAS
# ══════════════════════════════════════════════════════════════════════════
h1(doc, "10. ASSINATURAS")
para(doc, f"Teresina (PI), {date.today().strftime('%d de %B de %Y').replace('January','janeiro').replace('February','fevereiro').replace('March','março').replace('April','abril').replace('May','maio').replace('June','junho').replace('July','julho').replace('August','agosto').replace('September','setembro').replace('October','outubro').replace('November','novembro').replace('December','dezembro')}.")

doc.add_paragraph()
doc.add_paragraph()

# Tabela de assinaturas (3 campos em branco para preenchimento)
t_ass = doc.add_table(rows=2, cols=3)
t_ass.style = 'Table Grid'
for ci, label in enumerate(["Elaboração", "Revisão técnica", "Aprovação / Ciência"]):
    cell_top = t_ass.rows[0].cells[ci]
    cell_top.text = ""
    set_cell_bg(cell_top, 'FFFFFF')
    cell_bot = t_ass.rows[1].cells[ci]
    cell_bot.text = label
    set_cell_bg(cell_bot, 'F1F5F9')
    for run in cell_bot.paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = CINZA_MED
    for row in t_ass.rows:
        row.cells[ci].height = Cm(2.0) if row == t_ass.rows[0] else Cm(0.8)
for ci in range(3):
    set_col_width(t_ass, ci, Cm(5.5))

doc.add_paragraph()
note(doc,
    "Preencher nome completo, matrícula e cargo na linha de elaboração. "
    "Revisor e aprovador a critério da chefia imediata.")

doc.add_paragraph()
doc.add_paragraph()

# Rodapé
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    f"NT-GCGEO-001/2026  ·  Gerado automaticamente em {date.today().strftime('%d/%m/%Y')}  ·  "
    "Pipeline v2 — 9/9 testes OK  ·  GCGEO / SEMARH-PI"
)
run.font.size = Pt(8)
run.font.color.rgb = CINZA_MED
run.font.italic = True

# ── Salvar ────────────────────────────────────────────────────────────────
output = OUT / "NT-GCGEO-001-2026_Desmatamento_PI.docx"
doc.save(str(output))
print(f"Nota Técnica gerada: {output}")
print(f"Tamanho: {output.stat().st_size / 1024:.0f} KB")
